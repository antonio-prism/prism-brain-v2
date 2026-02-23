"""
PRISM Engine — AI-Powered Client Prefill

Uses Claude API with web search + optional uploaded documents to analyze a
company and suggest relevant business processes and risk events from the
PRISM catalogs. Returns structured JSON with process/risk selections and
vulnerability/resilience estimates.
"""

import base64
import json
import logging
import re
from functools import lru_cache
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)

# Paths to catalog data
SEEDS_DIR = Path(__file__).parent.parent / "frontend" / "data" / "seeds"
PROCESS_FRAMEWORK_PATH = Path(__file__).parent.parent / "frontend" / "data" / "process_framework.json"

# File upload limits
MAX_FILES = 5
MAX_TOTAL_SIZE_MB = 20
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt", ".csv"}

# Output JSON schema for Claude
OUTPUT_SCHEMA = {
    "type": "object",
    "properties": {
        "company_analysis": {
            "type": "string",
            "description": "2-3 paragraph analysis of the company based on web research and documents"
        },
        "processes": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "process_id": {"type": "string"},
                    "process_name": {"type": "string"},
                    "scope": {"type": "string"},
                    "rationale": {"type": "string"}
                },
                "required": ["process_id", "rationale"]
            }
        },
        "risks": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "event_id": {"type": "string"},
                    "event_name": {"type": "string"},
                    "domain": {"type": "string"},
                    "family_name": {"type": "string"},
                    "vulnerability": {"type": "number"},
                    "resilience": {"type": "number"},
                    "rationale": {"type": "string"}
                },
                "required": ["event_id", "vulnerability", "resilience", "rationale"]
            }
        }
    },
    "required": ["company_analysis", "processes", "risks"]
}


# ── Catalog Building ──────────────────────────────────────────────────

@lru_cache(maxsize=1)
def build_compact_catalogs() -> tuple:
    """Load and compress process framework + risk events into compact text.

    Returns: (process_text, risk_text, valid_process_ids_set, valid_event_ids_dict)
    - valid_process_ids_set: set of valid sub-process IDs (depth=2)
    - valid_event_ids_dict: dict mapping event_id -> {event_name, domain, family_name}
    """
    # Load process framework (only depth=2 sub-processes)
    process_lines = []
    valid_process_ids = set()
    process_lookup = {}

    with open(PROCESS_FRAMEWORK_PATH, "r", encoding="utf-8") as f:
        processes = json.load(f)

    for p in processes:
        if p.get("depth") == 2:
            pid = p["process_id"]
            valid_process_ids.add(pid)
            process_lookup[pid] = p
            process_lines.append(f"{pid}|{p['name']}|{p.get('scope', '')}|{p.get('scope_name', '')}")

    process_text = "\n".join(process_lines)

    # Load risk events from all 4 seed files
    risk_lines = []
    valid_event_ids = {}

    seed_files = sorted(SEEDS_DIR.glob("*_seed.json"))
    for seed_file in seed_files:
        with open(seed_file, "r", encoding="utf-8") as f:
            events = json.load(f)
        for e in events:
            eid = e["event_id"]
            valid_event_ids[eid] = {
                "event_name": e["event_name"],
                "domain": e["domain"],
                "family_name": e.get("family_name", ""),
                "family_code": e.get("family_code", ""),
                "base_rate_pct": e.get("base_rate_pct", 0),
            }
            desc = e.get("description", "")[:120]  # truncate long descriptions
            geo = e.get("geographic_scope", "")
            risk_lines.append(
                f"{eid}|{e['event_name']}|{e['domain']}|{e.get('family_name', '')}|{desc}|{geo}"
            )

    risk_text = "\n".join(risk_lines)

    logger.info(f"Built catalogs: {len(valid_process_ids)} processes, {len(valid_event_ids)} risk events")
    return process_text, risk_text, valid_process_ids, valid_event_ids


# ── Document Processing ───────────────────────────────────────────────

def process_uploaded_files(files: list) -> list:
    """Process uploaded files into content suitable for the Claude API.

    Args:
        files: List of FastAPI UploadFile objects

    Returns:
        List of dicts with keys: filename, type, content_block, info
        - PDF: content_block is a dict for Claude's native document block
        - Others: content_block is a dict with type="text" and text content
    """
    if not files:
        return []

    # Validate count
    if len(files) > MAX_FILES:
        raise ValueError(f"Maximum {MAX_FILES} files allowed, got {len(files)}")

    results = []
    total_size = 0

    for file in files:
        filename = file.filename or "unknown"
        ext = Path(filename).suffix.lower()

        # Validate extension
        if ext not in ALLOWED_EXTENSIONS:
            logger.warning(f"Skipping unsupported file type: {filename} ({ext})")
            continue

        # Read file content
        content = file.file.read()
        file.file.seek(0)  # Reset for potential re-read
        total_size += len(content)

        # Validate total size
        if total_size > MAX_TOTAL_SIZE_MB * 1024 * 1024:
            raise ValueError(f"Total file size exceeds {MAX_TOTAL_SIZE_MB}MB limit")

        try:
            if ext == ".pdf":
                result = _process_pdf(filename, content)
            elif ext == ".docx":
                result = _process_docx(filename, content)
            elif ext == ".xlsx":
                result = _process_xlsx(filename, content)
            else:  # .txt, .csv
                result = _process_text(filename, content)

            if result:
                results.append(result)
        except Exception as e:
            logger.warning(f"Could not process {filename}: {e}")
            results.append({
                "filename": filename,
                "type": ext.lstrip("."),
                "content_block": {"type": "text", "text": f"[Could not read {filename}: {e}]"},
                "info": {"error": str(e)},
            })

    return results


def _process_pdf(filename: str, content: bytes) -> dict:
    """Process PDF — send natively to Claude API as base64 document."""
    encoded = base64.standard_b64encode(content).decode("utf-8")
    # Estimate pages (rough: ~3KB per page for typical business PDFs)
    est_pages = max(1, len(content) // 3000)
    return {
        "filename": filename,
        "type": "pdf",
        "content_block": {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": encoded,
            },
        },
        "info": {"pages": est_pages, "size_bytes": len(content)},
    }


def _process_docx(filename: str, content: bytes) -> dict:
    """Process Word document — extract text via python-docx."""
    import io
    from docx import Document

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            text += "\n\n" + "\n".join(rows)

    return {
        "filename": filename,
        "type": "docx",
        "content_block": {
            "type": "text",
            "text": f"### Document: {filename}\n\n{text[:100000]}",  # cap at 100K chars
        },
        "info": {"paragraphs": len(paragraphs), "tables": len(doc.tables)},
    }


def _process_xlsx(filename: str, content: bytes) -> dict:
    """Process Excel file — extract sheets as markdown tables via openpyxl."""
    import io
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    text_parts = [f"### Spreadsheet: {filename}"]
    sheet_count = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):  # skip empty rows
                rows.append(" | ".join(cells))
        if rows:
            sheet_count += 1
            text_parts.append(f"\n**Sheet: {sheet_name}** ({len(rows)} rows)")
            # Include header + first 200 rows
            for r in rows[:201]:
                text_parts.append(r)
            if len(rows) > 201:
                text_parts.append(f"... ({len(rows) - 201} more rows)")

    wb.close()
    text = "\n".join(text_parts)

    return {
        "filename": filename,
        "type": "xlsx",
        "content_block": {
            "type": "text",
            "text": text[:100000],  # cap at 100K chars
        },
        "info": {"sheets": sheet_count},
    }


def _process_text(filename: str, content: bytes) -> dict:
    """Process plain text or CSV file."""
    # Try UTF-8 first, then latin-1
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1")

    return {
        "filename": filename,
        "type": Path(filename).suffix.lstrip("."),
        "content_block": {
            "type": "text",
            "text": f"### Document: {filename}\n\n{text[:100000]}",  # cap at 100K chars
        },
        "info": {"chars": len(text)},
    }


# ── Prompt Building ───────────────────────────────────────────────────

SYSTEM_PROMPT = """You are PRISM Brain's risk intelligence assistant. You research companies \
and select relevant business processes and risk events from predefined catalogs.

You have access to web search. Use it to research the company before making \
selections. Search for: company operations, industry position, geographic \
presence, supply chain structure, recent incidents, and known risk factors.

If the user has provided company documents (annual reports, risk registers, \
audit reports, etc.), review them carefully FIRST. These contain the most \
reliable, company-specific information. Use web search to supplement and \
verify what you find in the documents.

Return ONLY valid JSON matching the specified schema. Do not include markdown \
code fences or any text outside the JSON object."""


def build_prefill_prompt(client: dict, process_catalog: str, risk_catalog: str,
                         documents: Optional[list] = None) -> tuple:
    """Build the system prompt and user message content for the Claude API call.

    Args:
        client: Client profile dict with name, industry, location, etc.
        process_catalog: Compact process catalog text
        risk_catalog: Compact risk catalog text
        documents: Optional list of processed document dicts from process_uploaded_files()

    Returns:
        (system_prompt, user_content) where user_content is a list of content blocks
    """
    # Build client profile section
    profile_text = f"""## Company Profile
Name: {client.get('name', 'Unknown')}
Industry: {client.get('industry', 'Not specified')}
Location: {client.get('location', 'Not specified')}
Revenue: {client.get('revenue', 0)} {client.get('currency', 'EUR')}
Employees: {client.get('employees', 0)}
Sectors: {client.get('sectors', 'Not specified')}
Primary Markets: {client.get('primary_markets', 'Not specified')}
Export %: {client.get('export_percentage', 0)}
Notes: {client.get('notes', 'None')}"""

    # Build documents section
    if documents:
        doc_names = ", ".join(d["filename"] for d in documents)
        docs_intro = f"\n\n## Company Documents\nThe following {len(documents)} document(s) have been provided: {doc_names}\nReview these carefully before making selections.\n"
    else:
        docs_intro = "\n\n## Company Documents\nNo documents provided. Rely on web search.\n"

    # Instructions
    instructions = """
## Instructions
1. If company documents were provided, review them first to understand the company's operations, risk exposure, and existing assessments.
2. Search the web for this company to supplement document findings (or as primary research if no documents were provided).
3. Select 15-40 business processes most critical for this specific company from the Process Catalog below.
4. Select 20-60 risk events that represent the most relevant threats from the Risk Catalog below.
5. For each risk, suggest vulnerability (0.0-1.0) and resilience (0.0-1.0) based on the company's profile, documents, and industry.
6. Provide a 1-sentence rationale for each selection. Reference specific documents when the selection is based on uploaded content.

## Process Catalog (sub-processes)
Format: process_id|name|scope|scope_name
"""

    catalog_section = f"""{process_catalog}

## Risk Catalog (events)
Format: event_id|name|domain|family|description|geographic_scope
{risk_catalog}

## Output Schema
{json.dumps(OUTPUT_SCHEMA, indent=2)}

Return ONLY valid JSON. No markdown fences, no extra text."""

    # Build content blocks array
    user_content = []

    # Add text profile + docs intro
    user_content.append({"type": "text", "text": profile_text + docs_intro})

    # Add document content blocks (PDFs as native documents, others as text)
    if documents:
        for doc in documents:
            user_content.append(doc["content_block"])

    # Add instructions + catalogs
    user_content.append({"type": "text", "text": instructions + catalog_section})

    return SYSTEM_PROMPT, user_content


# ── Claude API Call ───────────────────────────────────────────────────

def call_claude_with_search(system_prompt: str, user_content: list,
                            api_key: str, retry: bool = True) -> dict:
    """Call Claude Sonnet with web search tool enabled.

    Args:
        system_prompt: System message
        user_content: List of content blocks for the user message
        api_key: Anthropic API key
        retry: Whether to retry on failure

    Returns:
        Parsed JSON dict from Claude's response
    """
    import anthropic

    client = anthropic.Anthropic(api_key=api_key)

    # Determine if documents are present (affects timeout)
    has_docs = any(
        block.get("type") == "document" or
        (block.get("type") == "text" and "### Document:" in block.get("text", ""))
        for block in user_content
    )
    # Longer timeout when processing documents
    timeout_seconds = 120 if has_docs else 90

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8192,
            system=system_prompt,
            tools=[{"type": "web_search_20250305"}],
            messages=[{"role": "user", "content": user_content}],
            timeout=timeout_seconds,
        )
    except Exception as e:
        logger.error(f"Claude API call failed: {e}")
        raise

    # Extract text content from response (may include tool_use blocks for web search)
    text_parts = []
    search_count = 0
    for block in response.content:
        if block.type == "text":
            text_parts.append(block.text)
        elif block.type == "web_search_tool_result":
            search_count += 1

    full_text = "\n".join(text_parts)

    # Parse JSON from response
    parsed = _extract_json(full_text)
    if parsed is None:
        if retry:
            logger.warning("Failed to parse JSON, retrying with explicit instruction")
            retry_content = user_content + [
                {"type": "text", "text": "\n\nYour previous response was not valid JSON. Please return ONLY a valid JSON object matching the schema."}
            ]
            return call_claude_with_search(system_prompt, retry_content, api_key, retry=False)
        raise ValueError(f"Could not parse JSON from Claude response. Raw text: {full_text[:500]}")

    parsed["_search_count"] = search_count
    parsed["_model"] = response.model
    return parsed


def _extract_json(text: str) -> Optional[dict]:
    """Extract JSON from Claude's response, handling markdown fences."""
    # Try direct parse first
    text = text.strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try removing markdown code fences
    match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1).strip())
        except json.JSONDecodeError:
            pass

    # Try finding the first { ... } block
    brace_start = text.find("{")
    if brace_start >= 0:
        # Find matching closing brace
        depth = 0
        for i in range(brace_start, len(text)):
            if text[i] == "{":
                depth += 1
            elif text[i] == "}":
                depth -= 1
                if depth == 0:
                    try:
                        return json.loads(text[brace_start:i + 1])
                    except json.JSONDecodeError:
                        break

    return None


# ── Response Validation ───────────────────────────────────────────────

def validate_ai_response(raw: dict, valid_process_ids: set,
                         valid_event_ids: dict) -> dict:
    """Validate and sanitize AI output.

    - Drops process_ids not in the framework
    - Drops event_ids not in the risk catalog
    - Clamps vulnerability/resilience to [0.05, 0.95]
    - Enriches with process_name/event_name from catalog lookups
    """
    # Validate processes
    valid_processes = []
    dropped_processes = 0
    for p in raw.get("processes", []):
        pid = str(p.get("process_id", ""))
        if pid in valid_process_ids:
            valid_processes.append(p)
        else:
            dropped_processes += 1
            logger.debug(f"Dropped invalid process_id: {pid}")

    if dropped_processes:
        logger.info(f"Dropped {dropped_processes} invalid process IDs from AI response")

    # Validate risks
    valid_risks = []
    dropped_risks = 0
    for r in raw.get("risks", []):
        eid = str(r.get("event_id", ""))
        if eid in valid_event_ids:
            # Clamp V/R values
            v = r.get("vulnerability", 0.5)
            res = r.get("resilience", 0.3)
            r["vulnerability"] = max(0.05, min(0.95, float(v)))
            r["resilience"] = max(0.05, min(0.95, float(res)))

            # Enrich with catalog data
            catalog_info = valid_event_ids[eid]
            r["event_name"] = catalog_info["event_name"]
            r["domain"] = catalog_info["domain"]
            r["family_name"] = catalog_info["family_name"]
            r["family_code"] = catalog_info["family_code"]
            r["base_rate_pct"] = catalog_info["base_rate_pct"]

            valid_risks.append(r)
        else:
            dropped_risks += 1
            logger.debug(f"Dropped invalid event_id: {eid}")

    if dropped_risks:
        logger.info(f"Dropped {dropped_risks} invalid event IDs from AI response")

    # Build validated result
    result = {
        "company_analysis": raw.get("company_analysis", "No analysis provided."),
        "processes": valid_processes,
        "risks": valid_risks,
        "validation": {
            "processes_suggested": len(raw.get("processes", [])),
            "processes_valid": len(valid_processes),
            "processes_dropped": dropped_processes,
            "risks_suggested": len(raw.get("risks", [])),
            "risks_valid": len(valid_risks),
            "risks_dropped": dropped_risks,
        }
    }

    # Low-confidence warnings
    warnings = []
    if len(valid_processes) < 5:
        warnings.append("AI returned fewer than 5 valid processes. Consider adding more manually.")
    if len(valid_risks) < 5:
        warnings.append("AI returned fewer than 5 valid risks. Consider adding more manually.")
    if warnings:
        result["warnings"] = warnings

    return result


# ── Main Orchestrator ─────────────────────────────────────────────────

def run_prefill(client: dict, api_key: str, files: Optional[list] = None) -> dict:
    """Run the full AI prefill pipeline.

    Args:
        client: Client profile dict
        api_key: Anthropic API key
        files: Optional list of FastAPI UploadFile objects

    Returns:
        Validated prefill result with processes, risks, and metadata
    """
    # 1. Build compact catalogs (cached)
    process_catalog, risk_catalog, valid_process_ids, valid_event_ids = build_compact_catalogs()

    # 2. Process uploaded documents
    documents = []
    documents_info = []
    if files:
        documents = process_uploaded_files(files)
        documents_info = [
            {"filename": d["filename"], "type": d["type"], **d.get("info", {})}
            for d in documents
        ]
        logger.info(f"Processed {len(documents)} documents for AI prefill")

    # 3. Build prompt
    system_prompt, user_content = build_prefill_prompt(
        client, process_catalog, risk_catalog, documents
    )

    # 4. Call Claude
    logger.info(f"Calling Claude for AI prefill: {client.get('name', 'unknown')} "
                f"({len(documents)} documents)")
    raw_response = call_claude_with_search(system_prompt, user_content, api_key)

    # 5. Validate
    result = validate_ai_response(raw_response, valid_process_ids, valid_event_ids)

    # 6. Add metadata
    result["documents_processed"] = documents_info
    result["model_used"] = raw_response.get("_model", "claude-sonnet-4-20250514")
    result["search_queries_used"] = raw_response.get("_search_count", 0)

    logger.info(
        f"AI prefill complete: {len(result['processes'])} processes, "
        f"{len(result['risks'])} risks, "
        f"{result['search_queries_used']} web searches"
    )

    return result
